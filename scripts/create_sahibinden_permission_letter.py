from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    os.environ.get(
        "SAHIBINDEN_PERMISSION_DOCX",
        "artifacts/Business_CEO_AI_Sahibinden_Sade_Izin_Onay_Yazisi.docx",
    )
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111827"
MUTED = "5F6B7A"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C8D1DC"
YELLOW = "FFF2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_paragraph_bottom_border(paragraph, color=BLUE, size=12, space=4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill=CALLOUT, border=BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def set_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(color)
    r_pr.append(sz)
    r.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    r.append(text)
    fld.append(r)
    paragraph._p.append(fld)


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_list_item(doc, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))


def add_label_value(doc, label: str, value: str, placeholder=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(f"{label}: "), bold=True)
    run = p.add_run(value)
    set_font(run)
    if placeholder:
        run.font.highlight_color = 7


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead) :]))
    else:
        set_font(p.add_run(text))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_font(hp.add_run("BUSINESS CEO AI"), size=9, bold=True, color=DARK_BLUE)
    hp.add_run("\t")
    set_font(
        hp.add_run("ERİŞİM İZİN TEYİDİ"),
        size=8.5,
        bold=True,
        color=MUTED,
    )

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)


def add_scope_table(doc: Document) -> None:
    headers = ["Erişim kapsamı", "Talep edilen veri ve kullanım", "Sahibinden kararı"]
    rows = [
        (
            "Arama sonuçları",
            "İlan numarası, ilan URL'si, başlık, konum, fiyat, ilan tarihi ve sayfalama bilgisinin filtreli sonuçlardan okunması.",
        ),
        (
            "İlan detayları",
            "Kamuya açık açıklama, gayrimenkul özellikleri ve ilan durumunun okunması; içerik üzerinde değişiklik yapılmaması.",
        ),
        (
            "Medya görüntüleme",
            "Kamuya açık fotoğraf URL'lerinin ve galeri sırasının işlenmesi. Yeniden yayımlama/kopyalama ayrıca onaylanmadıkça yapılmaz.",
        ),
        (
            "İletişim bilgileri",
            "Yalnızca açıkça izin verilen, kamuya gösterilen iletişim bilgisinin satış yetkisi görüşmesi amacıyla şifreli saklanması; önerilen azami saklama 90 gün.",
        ),
        (
            "Durum doğrulama",
            "İlanın aktif, kaldırılmış veya güncellenmiş olduğunun düşük sıklıkla kontrol edilmesi.",
        ),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 5520, 1740])
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), size=9.5, bold=True, color=DARK_BLUE)

    for scope, detail in rows:
        cells = table.add_row().cells
        set_table_geometry(table, [2100, 5520, 1740])
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_font(p0.add_run(scope), size=9.5, bold=True)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.05
        set_font(p1.add_run(detail), size=9.2)
        p2 = cells[2].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.05
        set_font(p2.add_run("☐ Onay\n☐ Şartlı\n☐ Red"), size=9.2)

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    set_font(
        after.add_run(
            "Not: İletişim bilgisi veya medya kapsamı onaylanmazsa sistem bu alanları toplamaz; kalan kapsam bağımsız değerlendirilebilir."
        ),
        size=9,
        italic=True,
        color=MUTED,
    )


def add_response_table(doc: Document) -> None:
    fields = [
        ("Karar", "☐ Onaylandı   ☐ Şartlı onaylandı   ☐ Onaylanmadı"),
        ("Yetkili erişim yöntemi", "☐ IP/WAF allowlist   ☐ API/veri akışı   ☐ Servis hesabı   ☐ Tarayıcı otomasyonu"),
        ("İzin verilen alan adları/uç noktalar", "........................................................................................................"),
        ("İzin verilen veri kapsamı", "........................................................................................................"),
        ("Hız/eşzamanlılık sınırı", "........................................................................................................"),
        ("İzin başlangıç ve bitiş tarihi", "........................................................................................................"),
        ("Sabit çıkış IP paylaşım kanalı", "........................................................................................................"),
        ("Ek koşullar", "........................................................................................................\n........................................................................................................"),
        ("Teknik irtibat", "Ad soyad: ........................................  E-posta: ........................................"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in fields:
        cells = table.add_row().cells
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_font(p0.add_run(label), size=9.5, bold=True, color=DARK_BLUE)
        set_cell_shading(cells[0], LIGHT_GRAY)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.05
        set_font(p1.add_run(value), size=9.4)
    set_table_geometry(table, [2700, 6660])


def add_signature_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run("sahibinden.com adına teyit eden"), bold=True, color=DARK_BLUE)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Ad soyad / unvan", "Kurumsal e-posta", "Tarih", "İmza / kaşe"]
    for row, label in zip(table.rows, labels):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_font(p0.add_run(label), size=9.5, bold=True, color=DARK_BLUE)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_font(p1.add_run("........................................................................................"), size=9.5)
    set_table_geometry(table, [2700, 6660])


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    bullet_num = add_numbering_definition(doc, bullet=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("OTOMATİK İLAN ERİŞİMİ"), size=23, bold=True, color="000000")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Kısa İzin ve Onay Teyidi"), size=14, color=MUTED)
    add_paragraph_bottom_border(subtitle)

    add_label_value(doc, "Muhatap", "sahibinden.com yetkili yöneticiliğine")
    add_label_value(doc, "Başvuru sahibi", "Business CEO AI - AI Portföy Uzmanı")
    add_label_value(doc, "Yasal şirket unvanı", "[KAŞEDE YER ALAN TAM UNVAN]", placeholder=True)
    add_label_value(doc, "Yetkili kişi", "[AD SOYAD / UNVAN]", placeholder=True)
    add_label_value(doc, "Kurumsal e-posta ve telefon", "[E-POSTA / TELEFON]", placeholder=True)
    add_label_value(doc, "Tarih", "11 Ağustos 2026")
    add_label_value(doc, "Konu", "Kamuya açık gayrimenkul ilanlarına otomatik erişim izninin yazılı teyidi")

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(12)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.line_spacing = 1.10
    shade_paragraph(callout)
    set_font(callout.add_run("Kısaca talebimiz. "), bold=True, color=DARK_BLUE)
    set_font(
        callout.add_run(
            "Daha önce tarafımıza bildirilen izni yazılı olarak teyit etmek ve uygulamanın hangi sınırlar içinde çalışabileceğini açıkça belirlemek istiyoruz."
        )
    )

    add_heading(doc, "Uygulama ne yapacak?", 1)
    add_body(
        doc,
        "Business CEO AI içindeki AI Portföy Uzmanı, kullanıcının seçtiği il, ilçe, mahalle ve gayrimenkul türüne göre sahibinden.com'daki kamuya açık gayrimenkul ilanlarını bulur. Bulunan ilanlar yalnızca ilgili kullanıcının kendi çalışma ekranında gösterilir ve satış yetkisi görüşmelerinin düzenli takip edilmesine yardımcı olur."
    )
    add_body(
        doc,
        "Uygulama ilan yayımlamaz, mevcut ilanları değiştirmez, kullanıcı hesabı adına işlem yapmaz ve kendiliğinden toplu mesaj göndermez."
    )

    add_heading(doc, "İzin talep ettiğimiz kapsam", 1)
    for item in (
        "Kamuya açık gayrimenkul ilanlarının uygulama tarafından arka planda otomatik olarak aranması ve okunması.",
        "İlan başlığı, konumu, fiyatı, açıklaması, özellikleri, fotoğrafları ve ilan bağlantısının kullanıcı ekranında gösterilmesi.",
        "İlanların yalnızca portföy araştırması ve satış yetkisi görüşmesi amacıyla kullanılması.",
        "Erişimin, sahibinden.com'un belirleyeceği kurallar ve özel koşullar çerçevesinde yürütülmesi.",
        "İznin geri çekilmesi hâlinde otomatik erişimin durdurulması.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "Sizden rica ettiğimiz teyit", 1)
    add_body(
        doc,
        "Yukarıda açıklanan kullanım için izin verilip verilmediğini aşağıdaki bölümde işaretlemenizi rica ederiz. Erişimin sağlanması için ayrıca yapılması gereken bir işlem varsa, teknik ayrıntıya girmenize gerek olmadan bizi yönlendireceğiniz kişi veya birimi belirtmeniz yeterlidir."
    )
    response = doc.add_table(rows=0, cols=2)
    response.style = "Table Grid"
    response_fields = [
        ("Karar", "☐ İzin veriyoruz   ☐ Şartlı izin veriyoruz   ☐ İzin vermiyoruz"),
        ("Varsa özel koşul", "........................................................................................\n........................................................................................"),
        ("Yönlendirilecek kişi/birim", "Ad soyad veya birim: ........................................................\nE-posta / telefon: .............................................................."),
        ("İzin süresi", "☐ Süresiz   ☐ İptal edilene kadar   ☐ Şu tarihe kadar: ......................"),
        ("Onaylayan yetkili", "Ad soyad / unvan: ..................................................................\nTarih / imza / kaşe: ................................................................"),
    ]
    for label, value in response_fields:
        cells = response.add_row().cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_font(p0.add_run(label), size=9.7, bold=True, color=DARK_BLUE)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.05
        set_font(p1.add_run(value), size=9.5)
    set_table_geometry(response, [2800, 6560])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(4)
    shade_paragraph(note, fill="F4F6F9", border=BLUE)
    set_font(note.add_run("Not: "), bold=True, color=DARK_BLUE)
    set_font(note.add_run("Erişim, sahibinden.com'un uygun gördüğü güvenli yöntemle sağlanacaktır. Bu yöntemin belirlenmesi için yukarıda bir kişi veya birim paylaşmanız yeterlidir."))

    add_heading(doc, "E-posta ile yanıt vermek isterseniz", 1)
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.22)
    quote.paragraph_format.right_indent = Inches(0.22)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(10)
    quote.paragraph_format.line_spacing = 1.10
    shade_paragraph(quote, fill=CALLOUT, border=BLUE)
    set_font(
        quote.add_run(
            "“sahibinden.com adına, Business CEO AI içindeki AI Portföy Uzmanı'nın kullanıcıların seçtiği ölçütlere göre kamuya açık gayrimenkul ilanlarını arka planda otomatik olarak aramasına ve bulunan ilanları yalnızca ilgili kullanıcının kendi çalışma ekranında göstermesine [İZİN VERİYORUZ / ŞARTLI İZİN VERİYORUZ / İZİN VERMİYORUZ]. Varsa özel koşullarımız: […]. Erişim konusunda iletişime geçilecek kişi veya birim: […].”"
        ),
        italic=True,
        color=DARK_BLUE,
    )

    add_heading(doc, "Başvuru sahibi beyanı", 1)
    add_body(
        doc,
        "Business CEO AI, sahibinden.com tarafından yazılı olarak bildirilen kapsam ve koşullara uyacağını; izin geri çekildiğinde otomatik erişimi durduracağını kabul eder."
    )
    sign = doc.add_paragraph()
    sign.paragraph_format.space_before = Pt(12)
    set_font(sign.add_run("Yetkili: "), bold=True)
    placeholder = sign.add_run("[AD SOYAD / UNVAN / İMZA / KAŞE]")
    set_font(placeholder)
    placeholder.font.highlight_color = 7

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.core_properties.title = "Business CEO AI - sahibinden.com Sade İzin Onay Yazısı"
    doc.core_properties.subject = "Kamuya açık gayrimenkul ilanlarına otomatik erişim izni teyidi"
    doc.core_properties.author = "Business CEO AI"
    doc.core_properties.keywords = "sahibinden.com, otomatik ilan erişimi, izin teyidi"
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
